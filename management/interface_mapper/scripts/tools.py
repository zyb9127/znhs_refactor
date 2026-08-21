"""
Interface Mapper Agent — 工具函数集

每个函数对应 SKILL.md 中定义的一个工具，供 Agent Loop 通过 Function Calling 调用。

工具列表：
  parse_docx       — 解析 docx 接口文档 → 结构化接口信息
  match_params     — 接口入参 × 主服务占位符 → request_template
  map_output       — 出参成功示例 → response_extract + field_transform
  detect_units     — 字段名/说明 → unit_convert 注入
  generate_skill   — 所有结果 → 写入 Skill 包目录文件
"""
from __future__ import annotations

import base64
import json
import os
import re
import uuid
from io import BytesIO
from pathlib import Path
from typing import Any, Dict, List, Optional

from loguru import logger

# ── 路径定位 ──────────────────────────────────────────────────
_MGMT_DIR   = Path(__file__).parent.parent            # management/interface_mapper/
_TMPL_DIR   = _MGMT_DIR / "template"
_CFG_DIR    = _MGMT_DIR / "config"
_PROJECT_ROOT = Path(__file__).parents[3]             # 项目根目录
_SKILLS_ROOT  = _PROJECT_ROOT / "skills-runtime"


def _load_json(path: Path) -> Dict:
    with open(path, encoding="utf-8") as f:
        return json.load(f)


def _service_params() -> Dict:
    return _load_json(_CFG_DIR / "main_service_params.json")


# ══════════════════════════════════════════════════════════════
# 工具 1：parse_docx
# ══════════════════════════════════════════════════════════════

def parse_docx(content_b64: str) -> Dict[str, Any]:
    """解析 base64 编码的 docx 文档，返回结构化接口信息。

    Args:
        content_b64: docx 文件的 base64 编码字符串

    Returns:
        {
          description, url, method, headers, version,
          input_params:  [{name, type, required, desc, example}],
          output_params: [{path, type, required, format, desc}],
          success_example: dict,
          fail_example: dict
        }
    """
    try:
        from docx import Document
    except ImportError:
        raise RuntimeError("请先安装 python-docx: pip install python-docx")

    raw = base64.b64decode(content_b64)
    doc = Document(BytesIO(raw))

    result: Dict[str, Any] = {
        "api_name": "",
        "description": "",
        "url": "",
        "method": "POST",
        "headers": {},
        "version": "",
        "input_params": [],
        "output_params": [],
        "success_example": {},
        "fail_example": {},
    }

    paragraphs = [p.text.strip() for p in doc.paragraphs]
    # 把表格内容也加入 full_text（表格内容不在 doc.paragraphs 里）
    table_texts = []
    for table in doc.tables:
        for row in table.rows:
            row_text = "\t".join(c.text.strip() for c in row.cells)
            if row_text.strip():
                table_texts.append(row_text)
    # 提取文本框/图文框内容（docx body 中的 txbx/wps:txbx 等 XML 元素）
    textbox_texts = _extract_textbox_texts(doc)
    full_text = "\n".join(paragraphs) + "\n" + "\n".join(table_texts) + "\n" + "\n".join(textbox_texts)

    # ── 通用字段提取（兼容同行/换行两种格式）──────────────────
    def _extract_field(text: str, *keywords: str) -> str:
        """尝试多种格式提取字段值：
        1. 同行冒号：关键词[：:]\s*值
        2. 同行 tab：关键词\t值（表格 key-value 行）
        3. 换行：关键词\n值
        """
        for kw in keywords:
            # 同行冒号格式：关键词：值 或 关键词:值
            m = re.search(rf"{kw}[：:]\s*([^\n\t]+)", text)
            if m:
                v = m.group(1).strip()
                if v:
                    return v
            # 同行 tab 格式（表格内容）：关键词\t值
            m = re.search(rf"^{kw}\t([^\n\t]+)", text, re.MULTILINE)
            if m:
                v = m.group(1).strip()
                if v:
                    return v
            # 换行格式：关键词\n值
            m = re.search(rf"{kw}[^\n]*\n\s*([^\n]+)", text)
            if m:
                v = m.group(1).strip()
                if v:
                    return v
        return ""

    # ── 同时扫描表格中的 key-value 行 ────────────────────────
    kv_from_tables: Dict[str, str] = {}
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if len(cells) >= 2 and cells[0] and cells[1]:
                kv_from_tables[cells[0]] = cells[1]

    def _kv_get(kv: Dict[str, str], *keywords: str) -> str:
        for kw in keywords:
            for k, v in kv.items():
                if kw in k:
                    return v.strip()
        return ""

    # ── 提取接口名称（优先从第一个表格的「接口名称」字段提取） ─────
    api_name_val = None

    # 优先策略：扫描所有表格，找第一张包含「接口名称」键的表格
    for table_idx, table in enumerate(doc.tables):
        if not table.rows:
            continue
        headers_row = [c.text.strip() for c in table.rows[0].cells]
        # 如果是键值对表格（常见于接口规范文档第一张表）
        if any("接口名称" in h or "接口名" in h for h in headers_row):
            for row in table.rows[1:]:
                cells = [c.text.strip() for c in row.cells]
                if len(cells) >= 2 and ("接口名称" in cells[0] or "接口名" in cells[0]):
                    api_name_val = cells[1] if len(cells) > 1 else cells[0]
                    logger.info(f"[parse_docx] 从第{table_idx+1}个表格的「接口名称」字段提取到: {api_name_val}")
                    break
        if api_name_val:
            break

    if not api_name_val:
        api_name_val = (
            _extract_field(full_text, "接口名称", "接口名", "API名称", "名称", "接口", "Api Name", "Name", "接口标题")
            or _kv_get(kv_from_tables, "接口名称", "接口名", "API名称", "名称", "接口", "Api Name")
        )

    # 最终兜底：文档第一个有意义的段落
    if not api_name_val and paragraphs:
        for p in paragraphs:
            p_clean = p.strip()
            if p_clean and len(p_clean) > 3 and not any(k in p_clean.lower() for k in ["参数", "字段", "说明", "输入", "输出", "示例", "版本", "请求", "地址"]):
                api_name_val = p_clean
                break

    if api_name_val:
        result["api_name"] = api_name_val.strip()
        logger.info(f"[parse_docx] ✅ 成功提取接口名称: {result['api_name']}")
    else:
        result["api_name"] = "未识别接口"
        logger.warning("[parse_docx] ⚠️ 未提取到接口名称，使用默认值 '未识别接口'")

    # ── 提取接口描述 ──────────────────────────────────────────
    desc = _extract_field(full_text, "接口描述", "功能描述", "描述") \
           or _kv_get(kv_from_tables, "接口描述", "功能描述", "描述")
    if desc:
        result["description"] = desc

    # ── 提取请求地址 ──────────────────────────────────────────
    url = _extract_field(full_text, "请求地址", "接口地址", "URL", "url") \
          or _kv_get(kv_from_tables, "请求地址", "接口地址", "URL")
    if url:
        result["url"] = url

    # ── 提取请求方式 ──────────────────────────────────────────
    method = _extract_field(full_text, "请求方式", "请求方法", "Method") \
             or _kv_get(kv_from_tables, "请求方式", "请求方法", "Method")
    if method:
        result["method"] = method.strip().upper().split()[0]  # 取第一个词，如 "POST application/json" → "POST"

    # ── 提取接口版本 ──────────────────────────────────────────
    ver = _extract_field(full_text, "接口版本", "版本") \
          or _kv_get(kv_from_tables, "接口版本", "版本")
    if ver:
        result["version"] = ver

    # ── 提取请求头（支持多个 header）────────────────────────────
    # 策略1：遍历 kv_from_tables，识别 HTTP header 格式的键值对
    _HTTP_HEADER_KEYWORDS = (
        "content-type", "x-channel-id", "x-app-id", "authorization",
        "accept", "x-request-id", "x-trace-id", "token", "appkey",
    )
    for k, v in kv_from_tables.items():
        k_lower = k.lower().strip()
        # 如果键名本身就是 HTTP header（含连字符 or 在已知列表中）
        if "-" in k_lower or any(kw in k_lower for kw in _HTTP_HEADER_KEYWORDS):
            if k.strip() and v.strip():
                result["headers"][k.strip()] = v.strip()
                logger.info(f"[parse_docx] 从表格 kv 中提取请求头: {k.strip()} = {v.strip()}")

    # 策略2：从文本段落中提取「请求头」区块，按行解析多个 Key: Value
    _header_section = ""
    for kw in ("请求头", "请求Header", "Request Header"):
        idx = full_text.find(kw)
        if idx >= 0:
            _header_section = full_text[idx: idx + 500]
            break
    if _header_section:
        for line in _header_section.splitlines():
            line = line.strip()
            if not line:
                continue
            # 匹配 "Key: Value" 或 "Key：Value"
            m = re.match(r'^([\w\-]+)\s*[：:]\s*(.+)$', line)
            if m:
                hk, hv = m.group(1).strip(), m.group(2).strip()
                # 排除章节标题行本身
                if hk not in ("请求头", "请求Header", "Request Header"):
                    result["headers"][hk] = hv
                    logger.info(f"[parse_docx] 从文本行中提取请求头: {hk} = {hv}")

    # 策略3：全文正则兜底，扫描常见 header 字段
    _COMMON_HEADERS_RE = re.compile(
        r'\b(Content-Type|X-Channel-Id|X-App-Id|Authorization|Accept|AppKey|Token'
        r'|X-Request-Id|X-Trace-Id)\s*[：:]\s*([^\n\t,;）)]+)',
        re.IGNORECASE,
    )
    for m in _COMMON_HEADERS_RE.finditer(full_text):
        hk = m.group(1).strip()
        hv = m.group(2).strip()
        # 只补充尚未提取到的 header（大小写不敏感去重）
        existing_lower = {k.lower() for k in result["headers"]}
        if hk.lower() not in existing_lower and hv:
            result["headers"][hk] = hv
            logger.info(f"[parse_docx] 全文正则兜底提取请求头: {hk} = {hv}")

    # 兜底：确保至少有 Content-Type
    if not any("content-type" == k.lower() for k in result["headers"]):
        if "application/json" in full_text.lower():
            result["headers"]["Content-Type"] = "application/json"

    # ── 解析表格（入参/出参）────────────────────────────────
    tables = doc.tables
    in_params: List[Dict] = []
    out_params: List[Dict] = []

    def _cell_text(cell) -> str:
        """合并单元格内所有段落文本（处理单元格内换行），去除首尾空白。"""
        return " ".join(p.text.strip() for p in cell.paragraphs if p.text.strip())

    for table in tables:
        rows = table.rows
        if not rows:
            continue
        headers_row = [_cell_text(c) for c in rows[0].cells]

        # 入参表：含"参数名" + 必选相关列（宽松匹配）
        _has_param_name = "参数名" in headers_row
        _has_required_col = any(k in headers_row for k in ("是否必选", "必选", "必填", "Required", "是否必填"))
        _has_type_col = any(k in headers_row for k in ("类型", "Type"))
        _has_out_col = any(k in headers_row for k in ("是否必返", "必返", "参数路径"))
        _is_input_table = _has_param_name and (_has_required_col or (_has_type_col and not _has_out_col))
        if _is_input_table:
            name_idx    = _col_idx(headers_row, ["参数名", "参数"])
            type_idx    = _col_idx(headers_row, ["类型", "Type"])
            req_idx     = _col_idx(headers_row, ["是否必选", "必选", "必填"])
            desc_idx    = _col_idx(headers_row, ["说明", "描述", "备注"])
            example_idx = _col_idx(headers_row, ["示例", "示例值", "样例"])

            for row in rows[1:]:
                cells = [_cell_text(c) for c in row.cells]
                if not any(cells):
                    continue
                name = cells[name_idx] if name_idx >= 0 else ""
                if not name or name.startswith("#"):
                    continue
                in_params.append({
                    "name":     name,
                    "type":     cells[type_idx]    if type_idx    >= 0 else "string",
                    "required": _is_required(cells[req_idx] if req_idx >= 0 else ""),
                    "desc":     cells[desc_idx]    if desc_idx    >= 0 else "",
                    "example":  cells[example_idx] if example_idx >= 0 else "",
                })

        # 出参表：含"参数路径"或"参数名"+"是否必返"
        elif ("参数路径" in headers_row or "参数名" in headers_row) and \
             ("是否必返" in headers_row or "必返" in headers_row or "说明" in headers_row):
            path_idx  = _col_idx(headers_row, ["参数路径", "参数名", "参数"])
            type_idx  = _col_idx(headers_row, ["类型", "Type"])
            req_idx   = _col_idx(headers_row, ["是否必返", "必返"])
            fmt_idx   = _col_idx(headers_row, ["格式", "Format"])
            desc_idx  = _col_idx(headers_row, ["说明", "描述"])

            for row in rows[1:]:
                cells = [_cell_text(c) for c in row.cells]
                if not any(cells):
                    continue
                path = cells[path_idx] if path_idx >= 0 else ""
                if not path or path.startswith("#"):
                    continue
                out_params.append({
                    "path":     path,
                    "type":     cells[type_idx] if type_idx >= 0 else "string",
                    "required": _is_required(cells[req_idx] if req_idx >= 0 else ""),
                    "format":   cells[fmt_idx]  if fmt_idx  >= 0 else "",
                    "desc":     cells[desc_idx] if desc_idx >= 0 else "",
                })

    result["input_params"]  = in_params
    result["output_params"] = out_params

    # ── 提取出参成功示例 JSON（多种常见标题，优先文档中的真实示例）────────────
    success_json: Optional[Dict] = None
    for _kw in (
        "出参成功示例",
        "成功响应示例",
        "响应报文示例",
        "返回示例",
        "响应示例",
    ):
        blk = _extract_json_block(full_text, before_keyword=_kw)
        if blk:
            success_json = blk
            break
    if not success_json and out_params:
        built = _build_mock_from_output_params(out_params)
        if built:
            success_json = built
            logger.info(
                "[parse_docx] 文档中未解析到出参成功示例 JSON，已根据出参表生成 mock 骨架"
            )
    if success_json:
        result["success_example"] = success_json

    # ── 提取出参失败示例 JSON ──────────────────────────────────
    fail_json = _extract_json_block(full_text, before_keyword="出参失败示例")
    if fail_json:
        result["fail_example"] = fail_json

    # ── 提取入参成功示例（用于入参匹配参考，兼容多种标题）──────────
    in_example_json = None
    for _in_kw in (
        "入参成功示例",
        "请求报文示例",
        "入参示例",
        "请求示例",
        "请求参数示例",
        "Request Body",
        "request body",
    ):
        in_example_json = _extract_json_block(full_text, before_keyword=_in_kw)
        if in_example_json:
            logger.info(f"[parse_docx] 从「{_in_kw}」中提取到入参示例: {list(in_example_json.keys())[:5]}")
            break
    if in_example_json:
        result["input_example"] = in_example_json
    else:
        logger.info("[parse_docx] 未找到入参成功示例 JSON，将使用入参表构建 request_template")

    # ── 调试日志：打印所有表格的表头，帮助诊断表格识别问题 ──────────
    for t_idx, table in enumerate(doc.tables):
        if table.rows:
            h = [c.text.strip() for c in table.rows[0].cells]
            logger.info(f"[parse_docx] 表格[{t_idx}] 表头: {h}  行数: {len(table.rows)}")

    logger.info(
        f"[parse_docx] 解析完成: url={result['url']} "
        f"input_params={len(in_params)} output_params={len(out_params)} "
        f"input_example={'有' if result.get('input_example') else '无'}"
    )
    return result


# ══════════════════════════════════════════════════════════════
# 工具 2：match_params
# ══════════════════════════════════════════════════════════════

def match_params(
    input_params_json: str,
    service_params_json: Optional[str] = None,
    input_example_json: Optional[str] = None,
) -> Dict[str, Any]:
    """将接口入参列表映射到主服务占位符，生成 request_template。

    当提供 input_example_json（入参成功示例 JSON）时，优先基于示例的真实结构递归构建
    request_template，同时自动检测 request_body_wrapper（如 params/body/data 等顶层包装键）。
    这样可以正确处理嵌套对象、避免带 "." 的字段名被当作字面量 key。

    Args:
        input_params_json:   接口入参列表 JSON 字符串
        service_params_json: 主服务占位符表 JSON（可选，默认加载配置文件）
        input_example_json:  入参成功示例 JSON 字符串（可选，优先使用）

    Returns:
        {
          "request_template":    {...},   # 可直接写入 api_nodes.json
          "request_body_wrapper": "...",  # 顶层包装键名，如 "params"（无则为 ""）
          "mapping_notes": [...]          # 匹配说明
        }
    """
    service_params = (
        json.loads(service_params_json)
        if service_params_json
        else _service_params()
    )
    placeholders = service_params.get("placeholders", {})

    # ── 优先路径：基于入参成功示例 JSON 递归构建 ──────────────────────────
    if input_example_json:
        try:
            example_data = (
                json.loads(input_example_json)
                if isinstance(input_example_json, str)
                else input_example_json
            )
            if isinstance(example_data, dict) and example_data:
                wrapper, template, notes = _build_template_from_example(
                    example_data, placeholders
                )
                logger.info(
                    f"[match_params] 基于入参示例构建完成: wrapper={wrapper!r} "
                    f"{len(template)} 个字段，{len(notes)} 个自定义"
                )
                return {
                    "request_template":     template,
                    "request_body_wrapper": wrapper,
                    "mapping_notes":        notes,
                }
        except Exception as e:
            logger.warning(f"[match_params] 入参示例解析失败，回退到入参表模式: {e}")

    # ── 兜底路径：基于扁平入参表构建（原有逻辑 + 自动检测 wrapper）──────
    input_params: List[Dict] = (
        json.loads(input_params_json)
        if isinstance(input_params_json, str)
        else input_params_json
    )

    template: Dict[str, Any] = {}
    notes: List[str] = []
    wrapper_key = ""

    # 检测是否存在单一 Object 顶层参数作为 wrapper（如 params/body/data）
    # 条件：顶层（无 "."）Object 类型参数只有一个，且名称在 WRAPPER_KEYS 中
    top_level_params = [
        p for p in input_params
        if p.get("name") and "." not in p["name"]
    ]
    if len(top_level_params) == 1:
        sole = top_level_params[0]
        if sole.get("type", "").lower() in ("object", "json对象") and \
                sole["name"].lower() in _WRAPPER_KEYS:
            wrapper_key = sole["name"]
            logger.info(f"[match_params] 入参表检测到 wrapper: {wrapper_key!r}，展开子字段")

    # 确定实际处理的入参集合（有 wrapper 时只处理其子字段，大小写不敏感匹配）
    if wrapper_key:
        prefix_lower = wrapper_key.lower() + "."
        prefix_len = len(wrapper_key) + 1   # wrapper_key 长度 + "." 的长度
        work_params = [
            {**p, "name": p["name"][prefix_len:]}  # 去掉 wrapper 前缀
            for p in input_params
            if p.get("name", "").lower().startswith(prefix_lower)
        ]
    else:
        work_params = input_params

    for param in work_params:
        name     = param.get("name", "")
        desc     = param.get("desc", "")
        typ      = param.get("type", "string").lower()
        example  = param.get("example", "")

        if not name:
            continue

        # 跳过带 "." 的路径形式字段名（已被父 object 处理，避免生成字面量 key）
        if "." in name:
            continue

        # 嵌套对象：如 crmpfPubInfo → 递归展开子字段
        # work_params 已剥离 wrapper 前缀，直接用 _build_nested_placeholder 即可
        if typ in ("object", "json对象") and not example:
            nested = _build_nested_placeholder(name, work_params, placeholders)
            if nested:
                template[name] = nested
            continue

        matched = _find_placeholder(name, desc, example, placeholders)
        if matched:
            template[name] = matched
        else:
            # extra_data 路径只用字段名本身（wrapper key 不属于 extra_data 路径）
            custom = "{{extra_data." + name + "}}"
            template[name] = custom
            notes.append(f"{name} → {custom}（从 extra_data 取值）")

    logger.info(f"[match_params] 入参匹配完成: wrapper={wrapper_key!r} {len(template)} 个字段，{len(notes)} 个自定义")
    return {"request_template": template, "request_body_wrapper": wrapper_key, "mapping_notes": notes}


# ── 入参示例递归构建辅助函数 ───────────────────────────────────────────

_WRAPPER_KEYS = {"params", "body", "data", "request", "req", "input", "param"}


def _build_template_from_example(
    example: Dict[str, Any],
    placeholders: Dict,
    _path_prefix: str = "",
) -> tuple:
    """递归遍历入参示例 JSON，为每个叶子字段匹配占位符，构建 request_template。

    Returns:
        (wrapper_key, template, notes)
        wrapper_key: 顶层包装键（如 "params"），无则为 ""
        template:    构建好的 request_template dict
        notes:       匹配说明列表
    """
    wrapper_key = ""

    # 检测顶层是否只有一个 wrapper key（如 {"params": {...}}）
    if not _path_prefix and len(example) == 1:
        only_key = next(iter(example))
        only_val = example[only_key]
        if isinstance(only_val, dict) and only_key.lower() in _WRAPPER_KEYS:
            wrapper_key = only_key
            example = only_val   # 展开 wrapper，以内部结构构建 template

    template, notes = _build_nested_from_example(example, placeholders, _path_prefix)
    return wrapper_key, template, notes


def _build_nested_from_example(
    data: Any,
    placeholders: Dict,
    path_prefix: str = "",
) -> tuple:
    """递归构建 template 节点。

    Returns:
        (node, notes)
    """
    if isinstance(data, dict):
        node: Dict[str, Any] = {}
        notes: List[str] = []
        for key, val in data.items():
            current_path = f"{path_prefix}.{key}" if path_prefix else key
            if isinstance(val, dict):
                child_node, child_notes = _build_nested_from_example(val, placeholders, current_path)
                node[key] = child_node
                notes.extend(child_notes)
            else:
                # 叶子节点：匹配占位符
                ph = _find_placeholder(key, "", str(val) if val is not None else "", placeholders)
                if not ph:
                    ph = "{{extra_data." + current_path + "}}"
                    notes.append(f"{current_path} → {ph}（从 extra_data 取值）")
                node[key] = ph
        return node, notes
    else:
        # 非 dict 叶子（不应出现在顶层，直接透传）
        return data, []


# ══════════════════════════════════════════════════════════════
# 工具 3：map_output
# ══════════════════════════════════════════════════════════════

# 主服务入参回显字段集合（出参映射时需跳过，不属于业务数据）
_SKIP_OUTPUT_FIELDS: set = {
    # 手机号
    "phone", "mobile", "msisdn", "phoneno", "phonenumber", "telno", "mobileno",
    # 意图
    "intent", "intentcode", "intentname",
    # 会话/追踪ID
    "callid", "sessionid", "traceid", "taskid", "requestid", "ioid", "logid", "serialno",
    # 省份
    "province", "botname", "provincecode", "areacode",
    # 其他
    "topn", "top", "pagesize", "pageno",
}


def _is_skip_field(field_name: str) -> bool:
    """判断字段是否为主服务入参回显字段（出参映射时应跳过）。
    比较时忽略大小写和下划线/中划线分隔符。
    """
    normalized = field_name.lower().replace("_", "").replace("-", "")
    return normalized in _SKIP_OUTPUT_FIELDS


def map_output(
    response_extract_json: str,
    field_transform_json: str,
    output_params_json: Optional[str] = None,
) -> Dict[str, Any]:
    """接收 LLM 分析出参文档后给出的映射结果，进行格式校验、空值过滤和规则校验。

    本工具不做任何自动探测或字段名匹配。LLM 需要：
    1. 阅读出参成功示例 JSON，确定实际根节点路径（不一定是 bean）
    2. 跳过主服务入参回显字段（phone/intent/callId/sessionId 等）
    3. 阅读出参说明，理解每个字段含义，映射到 7 大标准数据域
    4. 每个字段只映射到一个数据域，不得重复
    5. 处理平铺/拆分/合并等情况，在 field_transform 中声明转换规则
    6. 对需要单位转换的字段，在 field_transform 对应规则中注入 unit_convert

    Args:
        response_extract_json: LLM 给出的 response_extract JSON 字符串
            格式：{"数据域名": "实际取数路径", ...}
            示例：{"current_package": "data.mainOffer", "recommended_packages": "result.list"}
        field_transform_json:  LLM 给出的 field_transform JSON 字符串
            格式参考 生成api_node规范.txt，支持 passthrough/filter_include/filter_exclude
        output_params_json:    出参说明列表 JSON（可选，供 detect_units 后续使用）

    Returns:
        {
          "response_extract": {...},   # 过滤空值后的提取规则
          "field_transform":  {...},   # 过滤空值后的转换规则
          "output_params":    [...],   # 透传出参说明（供 detect_units 使用）
          "analysis": "...",
          "warnings": [...]            # 校验警告（重复字段、跳过字段等）
        }
    """
    response_extract: Dict[str, Any] = (
        json.loads(response_extract_json)
        if isinstance(response_extract_json, str)
        else response_extract_json
    )
    field_transform: Dict[str, Any] = (
        json.loads(field_transform_json)
        if isinstance(field_transform_json, str)
        else field_transform_json
    )
    out_params: List[Dict] = []
    if output_params_json:
        out_params = (
            json.loads(output_params_json)
            if isinstance(output_params_json, str)
            else output_params_json
        )

    warnings: List[str] = []

    # ── 过滤 response_extract 中的空值（NULL/""，保留 "0"）──────
    cleaned_extract: Dict[str, Any] = {
        k: v for k, v in response_extract.items()
        if v is not None and v != ""
    }

    # ── 过滤 field_transform 中的空值 ─────────────────────────
    cleaned_transform: Dict[str, Any] = {}
    for k, v in field_transform.items():
        if v is None or v == "":
            continue
        if isinstance(v, dict):
            # 过滤规则内部的空值，但保留结构
            cleaned_rule = {rk: rv for rk, rv in v.items()
                            if rv is not None and rv != ""}
            if cleaned_rule:
                cleaned_transform[k] = cleaned_rule
        else:
            cleaned_transform[k] = v

    # ── 校验1：检测 include_keys 中的主服务入参回显字段 ──────────
    for target_path, rule in cleaned_transform.items():
        if not isinstance(rule, dict):
            continue
        include_keys = rule.get("include_keys", [])
        for field in list(include_keys):
            if _is_skip_field(field):
                warnings.append(
                    f"[skip_field] {target_path}.include_keys 中的字段 '{field}' "
                    f"是主服务入参回显字段，已从映射中移除"
                )
                include_keys.remove(field)
                logger.warning(
                    f"[map_output] ⚠️ 跳过主服务入参字段: {target_path} → {field}"
                )
        if include_keys != rule.get("include_keys", []):
            cleaned_transform[target_path] = {**rule, "include_keys": include_keys}

    # ── 校验2：检测 include_keys 跨域重复字段 ────────────────────
    seen_fields: Dict[str, str] = {}   # field_name → first_target_path
    for target_path, rule in cleaned_transform.items():
        if not isinstance(rule, dict):
            continue
        include_keys = rule.get("include_keys", [])
        for field in include_keys:
            if field in seen_fields:
                warnings.append(
                    f"[duplicate_field] 字段 '{field}' 同时出现在 "
                    f"'{seen_fields[field]}' 和 '{target_path}' 的 include_keys 中，"
                    f"违反「同一字段只映射到一个数据域」规则"
                )
                logger.warning(
                    f"[map_output] ⚠️ 重复映射字段: '{field}' "
                    f"在 {seen_fields[field]} 和 {target_path} 中均出现"
                )
            else:
                seen_fields[field] = target_path

    if warnings:
        logger.warning(f"[map_output] 校验警告 ({len(warnings)} 条): {warnings}")

    analysis_parts = [
        f"response_extract: {list(cleaned_extract.keys())}",
        f"field_transform: {list(cleaned_transform.keys())}",
    ]
    if warnings:
        analysis_parts.append(f"warnings: {len(warnings)} 条")

    logger.info(
        f"[map_output] 映射完成: extract={list(cleaned_extract.keys())} "
        f"transform={list(cleaned_transform.keys())}"
    )
    return {
        "response_extract": cleaned_extract,
        "field_transform":  cleaned_transform,
        "output_params":    out_params,
        "analysis":         "；".join(analysis_parts),
        "warnings":         warnings,
    }


# ══════════════════════════════════════════════════════════════
# 工具 4：detect_units
# ══════════════════════════════════════════════════════════════

def detect_units(
    field_transform_json: str,
    output_params_json: Optional[str] = None,
) -> Dict[str, Any]:
    """扫描字段名/出参说明，推断单位转换需求，注入 unit_convert 规则。

    Args:
        field_transform_json: 当前 field_transform JSON 字符串
        output_params_json:   出参说明列表 JSON 字符串（可选，用于辅助识别）

    Returns:
        增强后的 field_transform（已注入 unit_convert 字段）
    """
    ft: Dict[str, Any] = (
        json.loads(field_transform_json)
        if isinstance(field_transform_json, str)
        else field_transform_json
    )
    out_params: List[Dict] = []
    if output_params_json:
        out_params = (
            json.loads(output_params_json)
            if isinstance(output_params_json, str)
            else output_params_json
        )

    # 构建 字段名 → 出参说明 映射（辅助识别）
    param_desc_map: Dict[str, str] = {}
    for p in out_params:
        path = p.get("path", "")
        desc = p.get("desc", "")
        key  = path.split(".")[-1] if path else ""
        if key:
            param_desc_map[key] = desc

    updated_ft = dict(ft)
    unit_conversions = []  # 新增：记录所有需要单位转换的字段，供前端展示

    for target_path, rule in ft.items():
        if target_path.startswith("_") or not isinstance(rule, dict):
            continue
        include_keys = rule.get("include_keys", [])
        if not include_keys:
            continue

        unit_convert: Dict[str, str] = {}
        field_rename: Dict[str, str] = {}  # 新增：仅当字段名含单位时生成重命名规则

        for field_name in include_keys:
            desc = param_desc_map.get(field_name, "")
            converter = _infer_converter(field_name, desc)
            if converter:
                unit_convert[field_name] = converter

                # 生成新字段名（仅在字段名包含原始单位标识时才重命名）
                new_name = _generate_renamed_field(field_name, converter)
                if new_name and new_name != field_name:
                    field_rename[field_name] = new_name

                unit_conversions.append({
                    "target_path": target_path,
                    "field": field_name,
                    "new_field": new_name or field_name,
                    "converter": converter,
                    "desc": desc
                })

        if unit_convert:
            updated_rule = dict(rule)
            updated_rule["unit_convert"] = unit_convert
            if field_rename:
                updated_rule["field_rename"] = field_rename
            updated_ft[target_path] = updated_rule
            logger.info(f"[detect_units] {target_path} 注入 unit_convert: {unit_convert} rename: {field_rename}")

    # 返回增强后的 field_transform + 单位转换列表（供前端展示）
    result = updated_ft
    result["_unit_conversions"] = unit_conversions  # 内部标记，前端可使用
    return result


# ══════════════════════════════════════════════════════════════
# 工具 5：generate_skill
# ══════════════════════════════════════════════════════════════

def generate_skill(
    province: str,
    intent: str,
    description: str,
    url: str,
    method: str,
    headers: Dict[str, str],
    request_template: Dict[str, Any],
    response_extract: Dict[str, Any],
    field_transform: Dict[str, Any],
    mock_response: Optional[Dict] = None,
    analysis: str = "",
    api_name: Optional[str] = None,
    api_display_name: Optional[str] = None,
    stage: str = "",
    scene: str = "",
    request_body_wrapper: str = "",
    field_aliases: Optional[Dict] = None,
    output_params: Optional[List[Dict]] = None,
) -> Dict[str, Any]:
    """渲染模板，生成并写入完整 Skill 包目录文件。

    Args:
        province:          省份代码（如 beijing）
        intent:            意图名称（如 套餐升档）
        description:       接口描述
        url:               接口 URL
        method:            请求方式
        headers:           请求头
        request_template:  入参模板
        response_extract:  响应提取规则
        field_transform:   字段转换规则
        mock_response:     Mock 响应（可选，从文档出参成功示例获取）
        analysis:          LLM 分析说明
        api_name:          接口节点名称（默认自动生成）
        api_display_name:  接口中文名称（来自文档"接口名称"字段，写入 _api_display_name）
        stage:             环节（用于接口节点命名前缀，如 upgrade）
        scene:             场景（用于接口节点命名前缀，如 inbound）
        field_aliases:     套餐字段别名（LLM 显式传入，优先级最高）
        output_params:     出参说明列表（来自 map_output 返回，用于自动推断省份专属字段别名）

    Returns:
        {
          "skill_path": "生成目录路径",
          "files": ["生成的文件列表"],
          "preview": {"api_nodes": {...}, "biz_config": {...}},
          "analysis": "..."
        }
    """
    skill_dir = _SKILLS_ROOT / province / intent
    skill_dir.mkdir(parents=True, exist_ok=True)
    (skill_dir / "config").mkdir(exist_ok=True)

    # 优先使用从文档解析出的真实接口名称（parse_docx 返回的 api_name / api_display_name）
    display_name = (api_display_name or api_name or "").strip()
    if not display_name:
        display_name = intent  # 兜底使用意图名称

    _api_name = api_name or _make_api_name(intent, stage=stage, scene=scene)

    # mock_response 为空时给出明确警告（应从文档 success_example 获取）
    if not mock_response:
        logger.warning(
            f"[generate_skill] ⚠️ mock_response 为空！接口 {_api_name} 在 mock_mode 下将返回 {{}}，"
            f"导致 response_extract 所有路径取值为 None，数据采集结果为空。"
            f"请确保将 parse_docx 返回的 success_example 作为 mock_response 传入。"
        )

    # ── 构造 api_nodes.json ────────────────────────────────────
    node_cfg: Dict[str, Any] = {
        "enabled":           True,
        "api_code":          f"{province}_{_api_name}_{str(uuid.uuid4())[:8]}",
        "_comment":          description,
        "_api_display_name": display_name,
        "api_display_name":  display_name,
        "api_name":          display_name,
        "url":               url,
        "method":            method,
        "headers":           headers,
        "timeout":           10,
        "max_retries":       2,
        "mock_mode":         True,
        "mock_response":     mock_response or {},
        "provide_domains":   _calc_provide_domains(response_extract, field_transform),
        "request_template":  request_template,
        "response_extract":  response_extract,
        "field_transform":   field_transform,
    }
    # 仅当 request_body_wrapper 非空时才写入（如北京的 "params"）
    if request_body_wrapper:
        node_cfg["request_body_wrapper"] = request_body_wrapper

    api_nodes: Dict[str, Any] = {
        "_comment":  f"{province} {intent} — 接口节点配置（由 interface_mapper Agent 自动生成）",
        "_desc":     "DataStep 并发调用本文件中 enabled=true 的所有接口",
        _api_name:   node_cfg,
    }

    # ── 构造 biz_config.json（默认模板，三路推断 field_aliases）──────────
    # 路1：从已构建的 api_nodes 扫描套餐域字段 + output_params 说明（最精准）
    inferred_from_nodes: Dict[str, List[str]] = {}
    try:
        inferred_from_nodes = _infer_field_aliases_from_api_nodes(api_nodes, output_params)
        if inferred_from_nodes:
            logger.info(f"[generate_skill] 从 api_nodes 推断 field_aliases: {inferred_from_nodes}")
    except Exception as e:
        logger.warning(f"[generate_skill] api_nodes 推断 field_aliases 失败（忽略）: {e}")

    # 路2：从 output_params 的路径说明推断（兼容无 field_transform 时的透传场景）
    inferred_from_params: Dict[str, List[str]] = {}
    if output_params and response_extract:
        try:
            inferred_from_params = _infer_field_aliases_from_output_params(output_params, response_extract)
            if inferred_from_params:
                logger.info(f"[generate_skill] 从 output_params 推断 field_aliases: {inferred_from_params}")
        except Exception as e:
            logger.warning(f"[generate_skill] output_params 推断 field_aliases 失败（忽略）: {e}")

    # 合并两路推断（api_nodes 推断优先，output_params 推断补充）
    combined_inferred: Dict[str, List[str]] = dict(inferred_from_params)
    for sem_key, fields in inferred_from_nodes.items():
        if sem_key not in combined_inferred:
            combined_inferred[sem_key] = []
        for f in fields:
            if f not in combined_inferred[sem_key]:
                combined_inferred[sem_key].insert(0, f)  # api_nodes 推断排在最前

    # 路3：LLM 显式传入 field_aliases（最高优先级，在 _merge_field_aliases 中处理）
    biz_config = _render_biz_config(
        province,
        intent,
        field_aliases=field_aliases,
        _inferred=combined_inferred,
    )

    # ── 构造 _meta.json ────────────────────────────────────────
    meta = {
        "name":              f"{province}-{intent}",
        "description":       description or f"{province} {intent} Skill",
        "version":           "1.0.0",
        "type":              "scenario",
        "province":          province,
        "scenario_id":       intent,
        "entry_script":      "scripts/main_flow.py",
        "config_files":      ["api_nodes.json", "biz_config.json"],
        "required_plugins":  ["plugins.package_diff", "plugins.unit_converter"],
        "author":            "interface_mapper_agent",
        "auto_generated":    True,
    }

    # ── 生成 SKILL.md ────────────────────────────────────────
    skill_md = _render_skill_md(province, intent, description, url, method,
                                api_nodes[_api_name])

    # ── 写入所有文件 ──────────────────────────────────────────
    files_written: List[str] = []

    def _write(path: Path, content: str):
        path.write_text(content, encoding="utf-8")
        files_written.append(str(path.relative_to(_PROJECT_ROOT)))

    _write(skill_dir / "config" / "api_nodes.json",
           json.dumps(api_nodes, ensure_ascii=False, indent=2))
    _write(skill_dir / "config" / "biz_config.json",
           json.dumps(biz_config, ensure_ascii=False, indent=2))
    _write(skill_dir / "_meta.json",
           json.dumps(meta, ensure_ascii=False, indent=2))
    _write(skill_dir / "SKILL.md", skill_md)

    logger.info(f"[generate_skill] ✅ Skill 包生成: {skill_dir}  文件={files_written}")
    return {
        "skill_path": str(skill_dir.relative_to(_PROJECT_ROOT)),
        "files":      files_written,
        "preview": {
            "api_nodes":       api_nodes,
            "biz_config":      biz_config,
            "success_example": mock_response or {},
        },
        "analysis": analysis,
    }


# ══════════════════════════════════════════════════════════════
# 私有工具函数
# ══════════════════════════════════════════════════════════════

def _col_idx(headers: List[str], candidates: List[str]) -> int:
    for c in candidates:
        if c in headers:
            return headers.index(c)
    return -1


def _is_required(val: str) -> bool:
    return "是" in val or "yes" in val.lower() or "true" in val.lower() or "1" == val.strip()


def _is_array_type(typ: str) -> bool:
    t = (typ or "").lower()
    return any(k in t for k in ("array", "list", "[]", "数组", "集合"))


def _normalize_path_segments(path: str) -> List[str]:
    """将「参数路径」拆成段，去掉 []、[0] 等记号。"""
    s = (path or "").strip().strip(".")
    if not s:
        return []
    s = re.sub(r"\[\d*\]", "", s)
    return [p for p in s.split(".") if p]


def _leaf_default_for_output(typ: str, leaf_name: str) -> Any:
    """根据出参类型与字段名生成占位值（用于无成功示例时的 mock 骨架）。"""
    leaf = (leaf_name or "").lower().replace("_", "")
    if leaf in ("resultcode", "code", "retcode", "rspcode"):
        return "0000"
    if leaf in ("resultmsg", "msg", "message", "retmsg", "rspdesc"):
        return "success"
    t = (typ or "string").lower()
    if _is_array_type(t):
        return []
    if any(x in t for x in ("int", "long", "number", "float", "double", "decimal")):
        return 0
    if "bool" in t:
        return False
    if "object" in t or "对象" in t:
        return {}
    return ""


def _collect_array_root_paths(rows: List[Dict]) -> set:
    """出参表中声明为数组类型的路径，以及 path 为 P 且存在多条 P.* 子路径时推断的数组根。"""
    roots: set = set()
    path_to_row: Dict[str, Dict] = {}
    stripped_paths: List[str] = []
    for p in rows:
        raw = (p.get("path") or "").strip()
        if not raw or raw.startswith("#"):
            continue
        key = raw.split("[")[0].strip(".").strip()
        path_to_row[key] = p
        stripped_paths.append(key)
        if _is_array_type(p.get("type", "")):
            roots.add(key)
    # 同一父路径下多条子路径 → 若父路径在表中单列为数组，已在上面覆盖；否则尝试推断
    children_by_parent: Dict[str, List[str]] = {}
    for key in stripped_paths:
        if "." not in key:
            continue
        parent, _ch = key.rsplit(".", 1)
        children_by_parent.setdefault(parent, []).append(key)
    for parent, kids in children_by_parent.items():
        if len(kids) < 2:
            continue
        prow = path_to_row.get(parent)
        if prow and _is_array_type(prow.get("type", "")):
            roots.add(parent)
    return roots


def _assign_mock_at_path(
    root: Dict[str, Any],
    segments: List[str],
    value: Any,
    array_roots: set,
) -> None:
    """按路径写入 mock，遇到声明为数组的路径段时使用 list 首元素 dict 导航。"""
    if not segments:
        return
    cur: Any = root
    for i, seg in enumerate(segments):
        is_last = i == len(segments) - 1
        acc = segments[: i + 1]
        prefix = ".".join(acc)
        if is_last:
            if not isinstance(cur, dict):
                return
            cur[seg] = value
            return
        if prefix in array_roots:
            if seg not in cur or not isinstance(cur[seg], list):
                cur[seg] = [{}]
            elif len(cur[seg]) == 0:
                cur[seg].append({})
            if not isinstance(cur[seg][0], dict):
                cur[seg][0] = {}
            cur = cur[seg][0]
        else:
            if seg not in cur or not isinstance(cur[seg], dict):
                cur[seg] = {}
            cur = cur[seg]


def _build_mock_from_output_params(output_params: List[Dict]) -> Dict[str, Any]:
    """接口文档无「出参成功示例」时，根据出参说明表生成最小 mock 结构，供 mock_response 使用。"""
    rows = [
        p
        for p in output_params
        if (p.get("path") or "").strip() and not (p.get("path") or "").startswith("#")
    ]
    if not rows:
        return {}
    array_roots = _collect_array_root_paths(rows)
    out: Dict[str, Any] = {}
    for p in rows:
        path = (p.get("path") or "").strip()
        typ = p.get("type", "string") or "string"
        segs = _normalize_path_segments(path)
        if not segs:
            continue
        leaf = segs[-1]
        val = _leaf_default_for_output(typ, leaf)
        _assign_mock_at_path(out, segs, val, array_roots)
    return out


def _extract_json_block(text: str, before_keyword: str = "") -> Optional[Dict]:
    """从文本中提取关键词之后第一个完整 JSON 块。

    增强策略：
    1. 先按关键词定位截断文本
    2. 找到第一个 '{' 后，逐字符匹配括号平衡，提取完整 JSON 块
    3. 多次尝试修复常见格式问题（多余换行、制表符等）
    """
    if before_keyword:
        idx = text.find(before_keyword)
        if idx >= 0:
            text = text[idx + len(before_keyword):]

    # 找到第一个 '{'
    start = text.find('{')
    if start < 0:
        return None

    # 括号平衡法提取完整 JSON 块（比贪婪正则更可靠）
    depth = 0
    in_str = False
    escape = False
    end = start
    for i, ch in enumerate(text[start:], start):
        if escape:
            escape = False
            continue
        if ch == '\\' and in_str:
            escape = True
            continue
        if ch == '"':
            in_str = not in_str
            continue
        if in_str:
            continue
        if ch == '{':
            depth += 1
        elif ch == '}':
            depth -= 1
            if depth == 0:
                end = i
                break

    candidate = text[start:end + 1]

    # 尝试直接解析
    try:
        return json.loads(candidate)
    except Exception:
        pass

    # 修复：去除多余空白/换行
    fixed = re.sub(r'[\r\n\t]+', ' ', candidate)
    fixed = re.sub(r'\s{2,}', ' ', fixed)
    try:
        return json.loads(fixed)
    except Exception:
        pass

    # 修复：用正则贪婪匹配兜底
    match = re.search(r'\{[\s\S]*\}', text)
    if match:
        try:
            return json.loads(match.group())
        except Exception:
            raw = re.sub(r'[\r\n]+', '\n', match.group())
            try:
                return json.loads(raw)
            except Exception:
                pass
    return None


def _get_path(data: Any, path: str) -> Any:
    if not path:
        return data
    v = data
    for key in path.split("."):
        if not isinstance(v, dict):
            return None
        v = v.get(key)
    return v


def _find_key_in_dict(
    d: Dict,
    candidates: List[str],
    expect_list: bool = False,
    expect_dict: bool = False,
) -> Optional[str]:
    for key in candidates:
        if key in d:
            val = d[key]
            if expect_list and not isinstance(val, list):
                continue
            if expect_dict and not isinstance(val, dict):
                continue
            return key
    return None


def _is_data_usage_key(key: str, desc: str = "") -> bool:
    """判断字段是否属于流量用量域（结合字段名和出参说明）"""
    combined = key + " " + desc
    kw = ["流量", "data", "MB", "GB", "流量饱和度", "流量(", "上网", "数据用量"]
    return any(w in combined for w in kw)


def _is_voice_usage_key(key: str, desc: str = "") -> bool:
    """判断字段是否属于语音用量域（结合字段名和出参说明）"""
    combined = key + " " + desc
    kw = ["语音", "主叫", "voice", "通话", "时长", "分钟", "语音用量"]
    return any(w in combined for w in kw)


def _is_consumption_key(key: str, desc: str = "") -> bool:
    """判断字段是否属于消费金额域（结合字段名和出参说明）"""
    combined = key + " " + desc
    kw = ["消费", "收入", "月均", "月消费", "折后", "费用", "arpu", "出账"]
    return any(w in combined for w in kw)


def _generate_renamed_field(field_name: str, converter: str) -> Optional[str]:
    """根据转换器生成重命名字段名，仅在字段名中包含原始单位标识时才替换。

    严格遵循用户要求：只有字段名称中含有 MB、分 等原始单位时才转换名称。
    示例：
      - "近3月平均流量(MB)" + mb_to_gb → "近3月平均流量(GB)"
      - "近3月平均月消费" + fen_to_yuan → None（无单位标识，不重命名）
      - "近6月平均折后收入（分）" + fen_to_yuan → "近6月平均折后收入（元）"
    """
    if not field_name or not converter:
        return None

    name = str(field_name).strip()

    if converter == "mb_to_gb":
        # 只在存在 MB 单位标识时才替换
        if re.search(r"MB|（MB）|\(MB\)|MB）|MB\s", name, re.IGNORECASE):
            name = re.sub(r"\(MB\)|（MB）|MB|MB）|MB\s", "(GB)", name, flags=re.IGNORECASE)
            name = re.sub(r"MB$", "(GB)", name, flags=re.IGNORECASE)
            return name
    elif converter == "fen_to_yuan":
        # 只在存在 分 单位标识时才替换
        if re.search(r"（分）|\(分\)|分|分）|（分）", name):
            name = re.sub(r"（分）|\(分\)|分|分）", "(元)", name)
            return name
    elif converter == "jiao_to_fen":
        if re.search(r"（角）|\(角\)|角|角）", name):
            name = re.sub(r"（角）|\(角\)|角|角）", "(分)", name)
            return name

    return None


def _infer_converter(field_name: str, desc: str = "") -> Optional[str]:
    """严格版：**仅当字段名称或接口文档说明中明确标注单位**时才返回转换器。

    符合用户最新要求：
    - 必须出现 MB、（MB）、(MB)、MB）、分、（分）、(分) 等明确单位标识
    - “流量饱和度”、“近3月流量”、“月消费”、“折后收入”等纯语义字段**不再**自动转换
    - 这会解决截图中“近3月流量饱和度”、“近3月流量”、“近3月平均月消费”等被错误标记的问题

    仅保留明确单位匹配，移除所有语义关键词列表。
    """
    if not field_name and not desc:
        return None

    combined = (field_name or "") + " " + (desc or "")

    # ── 严格仅明确单位标注才转换 ──────────────────────────────────────
    # MB 类（支持多种括号和位置）
    if re.search(r"MB|（MB）|\(MB\)|MB）|MB\s|单位[：:为]?[ \u00A0]*MB", combined, re.IGNORECASE):
        return "mb_to_gb"

    # 分 类（支持多种写法）
    if re.search(r"（分）|\(分\)|分[）)》\s$]|单位[：:为]?[ \u00A0]*分|单位为?分", combined):
        return "fen_to_yuan"

    # 角 类
    if re.search(r"（角）|\(角\)|角[）)》\s$]|单位[：:为]?[ \u00A0]*角|单位为?角", combined):
        return "jiao_to_fen"

    return None


# 主服务固定参数：字段名（小写无分隔符）→ 占位符（与 main_service_params.json 同步）
_FIXED_PARAM_MAP: Dict[str, str] = {
    "phone":    "{{PHONE}}",
    "mobile":   "{{PHONE}}",
    "msisdn":   "{{PHONE}}",
    "intent":   "{{INTENT}}",
    "callid":   "{{CALL_ID}}",
    "taskid":   "{{CALL_ID}}",
    "traceid":  "{{CALL_ID}}",
    "province": "{{PROVINCE}}",
    "botname":  "{{PROVINCE}}",
    "topn":     "{{TOP_N}}",
    "top":      "{{TOP_N}}",
    "curoffername": "{{CURRENT_OFFER_NAME}}",
    "curofferid":   "{{CURRENT_OFFER_ID}}",
    "curofferfee":  "{{CURRENT_OFFER_FEE}}",
}


def _find_placeholder(
    name: str, desc: str, example: str, placeholders: Dict
) -> Optional[str]:
    """将接口入参字段映射到占位符。

    规则（优先级从高到低）：
    1. 主服务固定参数（phone/intent/callId 等）→ {{PHONE}} 等固定占位符
    2. 手机号语义匹配 → {{PHONE}}
    3. 其他字段 → 返回 None，由调用方生成 {{extra_data.xxx}} 格式
    """
    name_lower = name.lower().replace("_", "").replace("-", "")

    # 1. 固定参数精确匹配
    if name_lower in _FIXED_PARAM_MAP:
        return _FIXED_PARAM_MAP[name_lower]

    # 2. 手机号语义匹配
    if "phone" in name_lower or "mobile" in name_lower or "手机" in desc:
        return "{{PHONE}}"

    # 3. 其他字段不预设占位符，由调用方生成 {{extra_data.xxx}}
    return None


def _build_nested_placeholder(
    parent: str,
    all_params: List[Dict],
    placeholders: Dict,
) -> Dict[str, Any]:
    """为嵌套对象构建占位符字典（子字段以 parent. 前缀筛选）

    子字段优先匹配固定参数，否则生成 {{extra_data.parent.child}} 格式。
    """
    obj: Dict[str, Any] = {}
    prefix = parent + "."
    for p in all_params:
        pname = p.get("name", "")
        if pname.startswith(prefix):
            child_name = pname[len(prefix):]
            ph = _find_placeholder(pname, p.get("desc", ""), p.get("example", ""), placeholders)
            obj[child_name] = ph or "{{extra_data." + pname + "}}"
    return obj if obj else {}


def _build_nested_placeholder_from_list(
    nested_name: str,
    child_params: List[Dict],
    placeholders: Dict,
    full_parent_path: str,
) -> Dict[str, Any]:
    """基于含原始完整路径的子字段列表，为嵌套对象递归构建占位符字典。

    Args:
        nested_name:      当前嵌套对象的字段名（如 crmpfPubInfo）
        child_params:     包含该嵌套对象子字段的入参列表（完整原始 name，如 params.crmpfPubInfo.staffId）
        placeholders:     主服务占位符字典
        full_parent_path: 完整的父路径（如 params.crmpfPubInfo），用于计算 extra_data 路径

    Returns:
        嵌套占位符字典，如 {"staffId": "{{extra_data.crmpfPubInfo.staffId}}"}
    """
    obj: Dict[str, Any] = {}
    prefix = full_parent_path + "."
    for p in child_params:
        pname = p.get("name", "")
        if not pname.startswith(prefix):
            continue
        # 取去掉完整父路径后的子字段名
        child_full = pname[len(prefix):]
        # 只处理直接子字段（不含 "."）
        if "." in child_full:
            continue
        desc = p.get("desc", "")
        ph = _find_placeholder(child_full, desc, p.get("example", ""), placeholders)
        if not ph:
            # extra_data 路径：去掉 wrapper 前缀（如 params.crmpfPubInfo.staffId → crmpfPubInfo.staffId）
            # full_parent_path 格式: "params.crmpfPubInfo"，去掉第一段
            parts = full_parent_path.split(".")
            extra_path = ".".join(parts[1:] + [child_full]) if len(parts) > 1 else f"{nested_name}.{child_full}"
            ph = "{{extra_data." + extra_path + "}}"
        obj[child_full] = ph
    return obj if obj else {}


def _make_api_name(intent: str, stage: str = "", scene: str = "") -> str:
    """从意图名称生成 api_nodes 中的接口节点名称。

    命名规范：{环节}_{场景}_{接口英文名}（全小写下划线）
    - 有 stage/scene 时拼接为前缀
    - 无 stage/scene 时直接用意图对应的英文名
    """
    mapping = {
        "套餐推荐": "package_recommend_api",
        "套餐升档": "package_upgrade_api",
        "流量包推荐": "flow_package_api",
    }
    base = mapping.get(intent) or (
        re.sub(r"[^\w\u4e00-\u9fff]", "_", intent).lower() + "_api"
    )
    parts = [p.lower().replace(" ", "_") for p in [stage, scene] if p and p.strip()]
    return "_".join(parts + [base]) if parts else base


# 默认字段别名（后备，各省接口字段名不同时在 field_aliases 中新字段名优先）
_DEFAULT_BIZ_FIELD_ALIASES: Dict[str, List[str]] = {
    "pkg_name":   ["offerName", "package_name", "productName", "name"],
    "pkg_fee":    ["initFee", "monthly_fee", "price", "fee"],
    "pkg_flow":   ["offerFlow", "data_quota", "dataGB", "flow"],
    "pkg_voice":  ["offerVoice", "voice_quota", "voiceMinutes", "voice"],
    "product_id": ["offerId", "product_id", "package_id", "offer_id"],
}

# field_aliases 推断：出参说明关键词 → 语义键
_ALIAS_SEMANTIC_RULES: List[tuple] = [
    ("pkg_name",   ["商品名称", "套餐名称", "产品名称", "方案名称", "offerName", "productName"]),
    ("pkg_fee",    ["月费", "价格", "费用", "月租", "套餐费", "fee", "price", "initFee"]),
    ("pkg_flow",   ["流量", "数据量", "GB", "流量额度", "flow", "offerFlow"]),
    ("pkg_voice",  ["语音", "通话分钟", "语音分钟", "分钟数", "voice", "offerVoice"]),
    ("product_id", ["商品标识", "产品ID", "套餐ID", "商品ID", "offerId", "productId", "packageId"]),
]


# 语义键推断：说明文字关键词 → 语义键（按优先级排列，先精确后宽泛）
_ALIAS_FROM_DESC_RULES: List[tuple] = [
    # product_id：标识类
    ("product_id", ["商品标识", "产品标识", "套餐标识", "套餐ID", "产品ID", "商品ID", "offerId", "productId"]),
    # pkg_name：名称类
    ("pkg_name",   ["商品名称", "套餐名称", "产品名称", "方案名称", "offerName", "productName"]),
    # pkg_fee：费用类
    ("pkg_fee",    ["月费", "月租", "套餐费", "价格", "费用", "fee", "price", "initFee"]),
    # pkg_flow：流量类
    ("pkg_flow",   ["流量", "数据量", "GB", "流量额度", "flow", "offerFlow"]),
    # pkg_voice：语音类
    ("pkg_voice",  ["语音", "通话分钟", "语音分钟", "分钟数", "voice", "offerVoice"]),
]

# 字段名模式推断（兜底，无出参说明时用字段名本身匹配）
_ALIAS_FROM_FIELD_NAME_RULES: List[tuple] = [
    ("product_id", [re.compile(r"(offerId|productId|packageId|offer_id|product_id|Id$)", re.I)]),
    ("pkg_name",   [re.compile(r"(offerName|productName|packageName|Name$)", re.I)]),
    ("pkg_fee",    [re.compile(r"(initFee|monthlyFee|price|Fee$|Price$|Cost$)", re.I)]),
    ("pkg_flow",   [re.compile(r"(offerFlow|dataQuota|dataGB|Flow$|Data$|GB$)", re.I)]),
    ("pkg_voice",  [re.compile(r"(offerVoice|voiceQuota|voiceMinutes|Voice$|Minute$)", re.I)]),
]


def _infer_alias_from_field(field_name: str, desc: str) -> Optional[str]:
    """对单个字段名+说明推断其语义键。

    优先级：说明关键词匹配 > 字段名模式匹配
    返回 None 表示无法推断（如描述类字段）。
    """
    combined_desc = desc.lower()
    combined_name = field_name.lower()

    # 1. 说明关键词匹配（说明有效时优先）
    if desc.strip():
        for sem_key, keywords in _ALIAS_FROM_DESC_RULES:
            if any(kw.lower() in combined_desc for kw in keywords):
                return sem_key

    # 2. 字段名模式匹配（说明为空或未匹配时兜底）
    for sem_key, patterns in _ALIAS_FROM_FIELD_NAME_RULES:
        for pat in patterns:
            if pat.search(field_name):
                return sem_key

    return None


def _collect_pkg_domain_fields(
    api_nodes: Dict[str, Any],
) -> Dict[str, str]:
    """从 api_nodes 中收集最终写入套餐域的字段名。

    扫描所有接口节点的 response_extract + field_transform，
    找出最终进入 current_package / recommended_packages 的字段名集合。

    Returns:
        {字段名: "来源描述"}（仅记录字段名，描述用于日志）
    """
    pkg_slots = {"current_package", "recommended_packages"}
    fields: Dict[str, str] = {}  # field_name → source hint

    for node_key, node in api_nodes.items():
        if node_key.startswith("_") or not isinstance(node, dict):
            continue

        extract = node.get("response_extract") or {}
        transform = node.get("field_transform") or {}
        mock = node.get("mock_response") or {}

        # 方式1：response_extract 直接命名为标准 slot 且 field_transform 无覆盖 → 从 mock 取字段
        for slot, path in extract.items():
            if slot not in pkg_slots:
                continue
            # 如果 field_transform 中对该 slot 有显式规则，跳过（下面单独处理）
            if slot in transform:
                continue
            # 直接透传：从 mock_response 按路径取值，枚举其字段名
            val = _get_path(mock, path) if path else None
            if isinstance(val, dict):
                for k in val:
                    if not k.startswith("_") and k not in fields:
                        fields[k] = f"{node_key}.{slot}(passthrough)"
            elif isinstance(val, list) and val and isinstance(val[0], dict):
                for k in val[0]:
                    if not k.startswith("_") and k not in fields:
                        fields[k] = f"{node_key}.recommended_packages[0](passthrough)"

        # 方式2：field_transform 中有 filter_include → include_keys 即为进入该域的字段
        for target_path, rule in transform.items():
            if target_path.startswith("_") or not isinstance(rule, dict):
                continue
            # target_path 的顶层 slot
            top_slot = target_path.split(".")[0]
            if top_slot not in pkg_slots:
                continue
            rule_type = rule.get("type", "passthrough")
            if rule_type in ("filter_include", "include"):
                for k in rule.get("include_keys", []):
                    if k not in fields:
                        fields[k] = f"{node_key}.{target_path}(include)"
            elif rule_type in ("passthrough",):
                # passthrough 整块透传：从 mock 按 from 字段路径取值，枚举字段名
                from_slot = rule.get("from", target_path)
                from_path = (extract or {}).get(from_slot, "")
                val = _get_path(mock, from_path) if from_path else None
                if isinstance(val, dict):
                    for k in val:
                        if not k.startswith("_") and k not in fields:
                            fields[k] = f"{node_key}.{target_path}(passthrough)"

    return fields


def _infer_field_aliases_from_api_nodes(
    api_nodes: Dict[str, Any],
    output_params: Optional[List[Dict]] = None,
) -> Dict[str, List[str]]:
    """从已构建的 api_nodes 推断 current_package / recommended_packages 域的字段别名。

    流程：
    1. 扫描 api_nodes，收集最终写入套餐域的字段名
    2. 查 output_params 的说明文字（按字段名叶子节点匹配）
    3. 用说明关键词或字段名模式推断语义键
    4. 返回 {sem_key: [字段名, ...]}

    Args:
        api_nodes:     generate_skill 构建好的接口节点 dict
        output_params: parse_docx 返回的出参说明列表（含 path 和 desc）

    Returns:
        {"pkg_name": ["curOfferName", ...], "pkg_fee": [...], ...}
    """
    # 构建 叶子字段名 → desc 的查找表
    leaf_to_desc: Dict[str, str] = {}
    for p in (output_params or []):
        path = (p.get("path") or "").strip()
        desc = (p.get("desc") or "").strip()
        if path and desc:
            leaf = path.split(".")[-1]
            if leaf and leaf not in leaf_to_desc:
                leaf_to_desc[leaf] = desc

    # 收集套餐域字段
    pkg_fields = _collect_pkg_domain_fields(api_nodes)
    if not pkg_fields:
        return {}

    result: Dict[str, List[str]] = {}
    for field_name, source in pkg_fields.items():
        desc = leaf_to_desc.get(field_name, "")
        sem_key = _infer_alias_from_field(field_name, desc)
        if sem_key:
            if sem_key not in result:
                result[sem_key] = []
            if field_name not in result[sem_key]:
                result[sem_key].append(field_name)
            logger.info(
                f"[field_aliases] {field_name!r}（{desc or '无说明'}）→ {sem_key}  来源={source}"
            )
        else:
            logger.debug(
                f"[field_aliases] {field_name!r}（{desc or '无说明'}）→ 无匹配语义键，跳过"
            )

    return result


def _infer_field_aliases_from_output_params(
    output_params: List[Dict],
    response_extract: Dict[str, Any],
) -> Dict[str, List[str]]:
    """从出参说明中推断 current_package / recommended_packages 域的字段别名。

    扫描属于套餐信息域（current_package / recommended_packages）的出参字段，
    结合字段路径和说明，映射到 pkg_name/pkg_fee/pkg_flow/pkg_voice/product_id 语义键。

    Args:
        output_params:    出参说明列表（parse_docx 返回）
        response_extract: 已确认的 response_extract 规则（用于判断哪些路径属于套餐域）

    Returns:
        {"pkg_name": ["curOfferName", ...], "pkg_fee": [...], ...}  仅含有推断结果的键
    """
    # 确定套餐域对应的路径前缀
    pkg_paths = set()
    for slot, path in response_extract.items():
        if slot in ("current_package", "recommended_packages") and path:
            # 取路径的父级前缀（如 bean.mainoffer → bean.mainoffer）
            pkg_paths.add(str(path).strip())

    result: Dict[str, List[str]] = {}

    for param in output_params:
        path = (param.get("path") or "").strip()
        desc = (param.get("desc") or "").strip()
        if not path:
            continue

        # 判断该字段是否属于套餐域路径下
        in_pkg_domain = any(path.startswith(pp) or path == pp for pp in pkg_paths)
        if not in_pkg_domain:
            continue

        # 取叶子字段名
        leaf = path.split(".")[-1].strip()
        if not leaf:
            continue

        # 按规则匹配语义键
        for sem_key, keywords in _ALIAS_SEMANTIC_RULES:
            combined = leaf + " " + desc
            if any(kw.lower() in combined.lower() for kw in keywords):
                if sem_key not in result:
                    result[sem_key] = []
                if leaf not in result[sem_key]:
                    result[sem_key].insert(0, leaf)  # 省份专属字段名排在最前
                break

    return result


def _merge_field_aliases(
    inferred: Dict[str, List[str]],
    explicit: Optional[Dict],
) -> Dict[str, Any]:
    """将推断别名与 LLM 显式传入的 field_aliases 合并，并与默认别名合并。

    优先级：explicit（LLM显式） > inferred（出参推断） > 默认兜底列表
    最终每个语义键的候选列表：[省份专属字段名...] + [默认通用字段名（去重）]
    """
    merged: Dict[str, List[str]] = {}

    # 以默认为基础
    for sem_key, defaults in _DEFAULT_BIZ_FIELD_ALIASES.items():
        base = list(inferred.get(sem_key, []))
        # 追加默认中不重复的字段名
        for d in defaults:
            if d not in base:
                base.append(d)
        merged[sem_key] = base

    # explicit 覆盖（LLM 显式传入优先级最高，插在列表最前）
    if isinstance(explicit, dict):
        for sem_key, val in explicit.items():
            if sem_key.startswith("_"):
                continue
            if isinstance(val, list):
                existing = merged.get(sem_key, [])
                new_list = list(val)
                for e in existing:
                    if e not in new_list:
                        new_list.append(e)
                merged[sem_key] = new_list

    return merged


def _render_biz_config(
    province: str,
    intent: str,
    field_aliases: Optional[Dict] = None,
    _inferred: Optional[Dict[str, List[str]]] = None,
) -> Dict:
    """生成 biz_config.json。

    field_aliases: LLM 调用 generate_skill 时显式传入的别名（优先级最高）
    _inferred:     由 generate_skill 已完成的两路推断结果（api_nodes 推断 + output_params 推断合并后）
    """
    final_aliases = _merge_field_aliases(_inferred or {}, field_aliases)

    return {
        "strategy": {
            "default_strategy": "direct",
            "top_n": 3,
            "intent_strategies": {},
            "max_script_length": 100,
            "max_parallel_scripts": 3,
            "expose_raw_bean": False,
        },
        "field_aliases": {
            "_comment": "套餐字段别名（按优先级列出，省份专属字段名在前）",
            **final_aliases,
        },
        "script_templates_v2": [
            {
                "template_id":       f"tpl_{province}_{intent}_default",
                "template_name":     f"{intent}话术模板",
                "intent":            intent,
                "province":          province,
                "product_id":        "",
                "scene":             intent,
                "stage":             intent,
                "template_content":  f"您好！根据您的套餐使用情况，为您推荐{{{{pkg_brief}}}}套餐，{'{diff_str}'}，更多实惠等您体验。",
                "linked_vars":       ["cur_brief", "pkg_brief", "diff_str", "usage_line"],
                "script_requirement": f"以用户专属客户经理的口吻自然、口语化沟通，别用官话套话；结合历史用量/标签点出最突出的痛点，再用推荐套餐真实字段值说清如何解决、做前后对比放大获得感；只讲有数据支撑的卖点，不编造不夸大；{intent}推荐话术贴合用户痛点，100字以内，结尾一句自然的办理引导。",
                "status":            "online",
                "created_by":        "interface_mapper_agent",
            }
        ],
    }


def _render_skill_md(
    province: str,
    intent: str,
    description: str,
    url: str,
    method: str,
    api_node_cfg: Dict,
) -> str:
    """生成省份 Skill 包的 SKILL.md"""
    extract_keys = list(api_node_cfg.get("response_extract", {}).keys())
    return f"""---
skill_id: {province}:{intent}
version: 1.0.0
province: {province}
intent: {intent}
entry: scripts/main_flow.py::run_scenario_flow
generated_by: interface_mapper_agent
input_contract:
  phone: "{{{{PHONE}}}} 用户手机号"
  intent: "{intent}"
  botName: "{province}"
output_contract:
  current_package: 用户当前套餐信息
  recommended_packages: 推荐套餐列表
  usage: 用户用量统计
  tags: 用户业务标签
  other_info: 接口原始响应数据
---

# {province.capitalize()} · {intent} Skill

## 描述
{description or f"{province} {intent} 智能推荐 Skill，由 interface_mapper Agent 自动生成"}

## 接口信息
| 项目 | 说明 |
|------|------|
| 请求地址 | `{url}` |
| 请求方式 | {method} |
| 接口描述 | {description} |

## 响应字段映射
提取的标准 Slot：{', '.join(extract_keys)}

## 执行流程
1. DataStep → 调用接口，两步 JSON 映射到 resource_context
2. RecommendStep → 按 direct 策略筛选 TopN 推荐产品
3. ScriptStep → 并发 LLM 生成个性化话术

## 入口
`scripts/main_flow.py::run_scenario_flow(context, request_data)`
"""


def _render_main_flow(province: str, intent: str) -> str:
    """生成标准省份 main_flow.py（通用模板，无硬编码 extra_vars）"""
    return f'''"""
{province.capitalize()} · {intent} — 省份编排入口（由 interface_mapper Agent 自动生成）

职责：
1. 构建 FlowContext（传入 extra_data，DataStep 自动解析占位符）
2. 调用三步管道
3. 透传原始 bean 数据到 other_info
"""
from __future__ import annotations
from typing import Any, Dict
from core.context import FlowContext
from core.pipeline import MarketingPipeline


async def run_scenario_flow(
    context: Any,               # SkillExecutionContext
    request_data: Dict[str, Any],
) -> Dict[str, Any]:
    """{province} {intent} 主流程"""
    ctx = FlowContext(
        phone=request_data.get("phone", ""),
        intent=request_data.get("intent", "{intent}"),
        province="{province}",
        top_n=int(request_data.get("topN", 3)),
        trace_id=getattr(context, "trace_id", ""),
        extra_data=request_data.get("extra_data") or {{}},
        extra_info=request_data.get("extra_info") or {{}},
        extra_context=request_data.get("extra_context") or {{}},
    )

    skill_config = getattr(context, "package", None)
    skill_config = skill_config.config if skill_config else {{}}

    result = await MarketingPipeline().execute(ctx, skill_config=skill_config)

    # 透传原始接口响应到 other_info
    first_raw = next(iter(ctx.raw_responses.values()), {{}})
    result["other_info"] = (
        first_raw.get("bean", first_raw) if isinstance(first_raw, dict) else first_raw
    )
    return result
'''


# 标准数据域白名单
_STD_SLOTS = {
    "current_package", "recommended_packages", "usage",
    "tags", "user_info", "user_profile", "domain_ext",
}


def _calc_provide_domains(
    response_extract: Dict[str, Any],
    field_transform: Dict[str, Any],
) -> List[str]:
    """从 response_extract + field_transform 推导该接口可提供的标准数据域列表。

    规则：
    - 扫描 response_extract 的 key（直接命名为标准 slot 的）
    - 扫描 field_transform 的 target_path（取顶层域名，如 usage.data_usage → usage）
    - 过滤掉以 _ 开头的注释 key，去重后排序
    """
    domains: set = set()
    for k in response_extract:
        if not k.startswith("_"):
            top = k.split(".")[0]
            if top in _STD_SLOTS:
                domains.add(top)
    for k in field_transform:
        if not k.startswith("_"):
            top = k.split(".")[0]
            if top in _STD_SLOTS:
                domains.add(top)
    return sorted(domains)


def _extract_textbox_texts(doc) -> List[str]:
    """提取 docx 文档中文本框（txbx/wps:txbx）内的所有文本段落。
    
    docx 文档中的文本框内容不在 doc.paragraphs 里，需要通过 XML 解析获取。
    支持标准 OOXML 文本框（w:txbx）和 WPS 兼容文本框（mc:AlternateContent）。
    """
    from lxml import etree
    
    texts: List[str] = []
    
    # XML 命名空间
    W_NS  = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    W14_NS = "http://schemas.microsoft.com/office/word/2010/wordml"
    
    try:
        body = doc.element.body
        
        # 递归提取所有 w:txbx 和 mc:AlternateContent 中的文本
        # 查找所有 w:txbx 节点
        txbx_nodes = body.findall(
            f".//{{{W_NS}}}txbx"
        )
        for txbx in txbx_nodes:
            # 提取节点下所有文本
            text = "".join(txbx.itertext())
            if text.strip():
                texts.append(text.strip())
                logger.info(f"[parse_docx] 从文本框(w:txbx)提取内容: {text.strip()[:100]}")
        
        # 如果没找到，尝试更宽泛的 XML 搜索
        if not txbx_nodes:
            xml_str = etree.tostring(body, encoding="unicode")
            # 查找所有 txbx 标签（忽略命名空间前缀）
            import re as _re
            # 提取文本框中的文本内容
            txbx_pattern = _re.compile(r'<w:txbx[^>]*>(.*?)</w:txbx>', _re.DOTALL)
            for m in txbx_pattern.finditer(xml_str):
                # 从 XML 片段中提取纯文本
                inner = m.group(1)
                t_pattern = _re.compile(r'<w:t[^>]*>(.*?)</w:t>', _re.DOTALL)
                t_texts = t_pattern.findall(inner)
                joined = "".join(t_texts)
                if joined.strip():
                    texts.append(joined.strip())
                    logger.info(f"[parse_docx] 从文本框(XML正则)提取内容: {joined.strip()[:100]}")
    except Exception as e:
        logger.warning(f"[parse_docx] 文本框提取失败（忽略）: {e}")
    
    return texts


