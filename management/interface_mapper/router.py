"""
接口映射路由模块

包含所有与接口文档解析、LLM 映射、迭代优化相关的路由：
  POST /api/parse_docx_preview       — 解析 docx 文档，返回完整预览（不写文件）
  POST /api/interfaces/auto_map      — 基于响应样例 LLM 自动生成映射规则
  POST /api/interfaces/apply_parsed  — 将解析结果写入 api_nodes.json
  POST /api/skills/preview_mapping   — 用 mock_response 跑映射，返回域结果预览
  POST /api/skills/refine_mapping_preview — LLM 反推规则（仅预览，不写文件）
  POST /api/skills/refine_mapping    — LLM 反推规则并写入 api_nodes.json
  POST /api/skills/create_from_doc   — Interface Mapper Agent 全自动生成 Skill 包
  GET  /api/download/interface_template — 下载接口规范文档模板
"""
from __future__ import annotations

import json
import re
from io import BytesIO
from pathlib import Path
from typing import Any, Dict, Optional

from fastapi import APIRouter, HTTPException
from loguru import logger
from pydantic import BaseModel, Field

from management.interface_mapper.scripts.prompts import (
    AUTO_MAP_SYSTEM_PROMPT,
    PARSE_DOCX_SYSTEM_PROMPT,
    REFINE_MAPPING_SYSTEM_PROMPT,
)
from management.interface_mapper.scripts.tools import detect_units

router = APIRouter(tags=["接口映射"])


# ── 依赖注入（延迟导入，避免循环）────────────────────────────────

def _get_llm():
    """获取接口映射专用 LLM 实例（dashscope）。"""
    from services.llm_service import LLMService
    from utils.config_loader import config_loader
    return LLMService(config_override=config_loader.get_dashscope_config())


def _get_skill_registry():
    from utils.skill_runtime import skill_registry
    return skill_registry


# ── 数据模型 ──────────────────────────────────────────────────

class ParseDocxPreviewRequest(BaseModel):
    province: str = Field(..., description="省份代码，如 beijing")
    intent: str = Field(..., description="意图名称，如 套餐升档")
    docx_content_b64: str = Field(..., description="docx 文件的 base64 编码内容")


class CreateSkillFromDocRequest(BaseModel):
    province: str = Field(..., description="省份代码，如 beijing")
    intent: str = Field(..., description="意图名称，如 套餐升档")
    docx_content_b64: str = Field(..., description="docx 文件的 base64 编码内容")


class PreviewMappingRequest(BaseModel):
    province: str
    intent: str
    mock_response: dict = Field(default_factory=dict)
    response_extract: dict = Field(default_factory=dict)
    field_transform: dict = Field(default_factory=dict)


class RefineMappingRequest(BaseModel):
    province: str
    intent: str
    user_domain_result: dict = Field(..., description="用户修改后的数据域结果")
    current_api_nodes: dict = Field(..., description="当前 api_nodes.json 内容")
    mock_response: dict = Field(default_factory=dict)


class RefineMappingPreviewRequest(BaseModel):
    """文档解析向导中反推规则：不读不写磁盘上的 api_nodes。"""
    mock_response: dict = Field(default_factory=dict)
    user_domain_result: dict = Field(..., description="用户修改后的数据域结果")
    response_extract: dict = Field(default_factory=dict)
    field_transform: dict = Field(default_factory=dict)


# ── 内部工具函数 ──────────────────────────────────────────────

def _apply_mapping(mock_response: dict, response_extract: dict, field_transform: dict) -> dict:
    """用 mock_response 跑 response_extract + field_transform，调用 DataStep 的映射方法。"""
    from steps.data_step import DataStep
    step = DataStep("preview")
    api_cfg = {"response_extract": response_extract, "field_transform": field_transform}
    extracted = step._extract_fields(mock_response, api_cfg)
    resources = step._transform_fields(extracted, api_cfg, mock_response)
    # 补充未经 transform 的 extract 域（临时 slot 如 raw_tags 也透出，方便调试）
    for slot_name, val in extracted.items():
        if slot_name not in resources:
            resources[slot_name] = val
    return resources


async def _run_refine_mapping_llm(
    mock_resp: dict,
    current_extract: dict,
    current_transform: dict,
    user_domain_result: dict,
    *,
    llm_stage: str = "interface_mapper.refine",
) -> tuple:
    """LLM 根据用户期望的数据域结果反推 response_extract / field_transform。

    Returns:
        (new_extract, new_transform, analysis)
    """
    prompt = REFINE_MAPPING_SYSTEM_PROMPT.format(
        mock_resp=json.dumps(mock_resp, ensure_ascii=False, indent=2),
        current_extract=json.dumps(current_extract, ensure_ascii=False, indent=2),
        current_transform=json.dumps(current_transform, ensure_ascii=False, indent=2),
        user_domain_result=json.dumps(user_domain_result, ensure_ascii=False, indent=2),
    )

    llm = _get_llm()
    llm_resp = await llm.generate(prompt, temperature=0.1, stage=llm_stage, max_tokens=2000)
    raw = (llm_resp or "").strip()
    m = re.search(r'\{[\s\S]*\}', raw)
    if not m:
        raise ValueError("LLM 返回内容无法解析为 JSON")
    refined = json.loads(m.group())

    new_extract = refined.get("response_extract", current_extract)
    new_transform = refined.get("field_transform", current_transform)
    analysis = refined.get("analysis", "")

    new_transform = detect_units(json.dumps(new_transform, ensure_ascii=False))
    return new_extract, new_transform, analysis


# ── 路由 ──────────────────────────────────────────────────────

@router.post("/api/interfaces/auto_map")
async def auto_map_interface(body: Dict[str, Any]):
    """基于 LLM 对接口响应样例自动生成 response_extract + field_transform 规则。

    body:
      sample_response: dict  接口响应样例
    """
    from steps.data_step import DataStep

    sample_response = body.get("sample_response")
    if not sample_response:
        raise HTTPException(400, "sample_response 不能为空")

    sample_str = json.dumps(sample_response, ensure_ascii=False, indent=2)
    if len(sample_str) > 4000:
        sample_str = sample_str[:4000] + "\n... (已截断)"

    prompt = AUTO_MAP_SYSTEM_PROMPT.replace("{sample_str}", sample_str)

    try:
        llm = _get_llm()
        raw_text = await llm.generate(
            prompt, temperature=0.1, max_tokens=8192,
            stage="auto_map", provider="auto_map",
        )
    except Exception as e:
        raise HTTPException(500, f"LLM 调用失败: {e}")

    json_match = re.search(r'\{[\s\S]*\}', raw_text)
    if not json_match:
        raise HTTPException(500, f"LLM 返回无法解析，原始内容: {raw_text[:300]}")
    try:
        mapping_cfg = json.loads(json_match.group())
    except Exception:
        raise HTTPException(500, f"LLM 返回 JSON 格式错误: {raw_text[:300]}")

    response_extract = mapping_cfg.get("response_extract", {})
    field_transform = mapping_cfg.get("field_transform", {})
    analysis = mapping_cfg.get("analysis", "")

    # detect_units 兜底：LLM 可能漏掉 unit_convert，规则引擎补扫一遍
    if field_transform:
        field_transform = detect_units(json.dumps(field_transform, ensure_ascii=False))

    # 应用生成的映射规则到样例，生成预览
    step = DataStep("auto_map")
    api_cfg_tmp = {"response_extract": response_extract, "field_transform": field_transform}
    extracted = step._extract_fields(sample_response, api_cfg_tmp)
    preview = step._transform_fields(extracted, api_cfg_tmp, sample_response)

    return {
        "code": 200,
        "data": {
            "response_extract": response_extract,
            "field_transform": field_transform,
            "analysis": analysis,
            "preview": preview,
        },
    }


@router.post("/api/parse_docx_preview")
async def parse_docx_preview(req: ParseDocxPreviewRequest):
    """解析接口文档并返回完整预览（不写文件）。

    工作流：
      Step 1  parse_docx   — 解析 docx，提取接口信息
      Step 2  match_params — 入参与 FlowContext 占位符匹配
      Step 3  LLM 辅助映射 — 分析出参字段，生成 response_extract + field_transform
      Step 4  detect_units — 注入单位转换规则（MB→GB，分→元）
    """
    try:
        from management.interface_mapper.scripts.tools import (
            parse_docx, match_params, detect_units as _detect_units,
        )

        # Step 1: 解析文档
        doc_info = parse_docx(req.docx_content_b64)

        # Step 2: 入参匹配（规则引擎，优先使用入参成功示例 JSON 智能推断结构）
        input_params_json = json.dumps(doc_info.get("input_params", []), ensure_ascii=False)
        input_example = doc_info.get("input_example")
        input_example_json = json.dumps(input_example, ensure_ascii=False) if input_example else None
        match_result = match_params(input_params_json, input_example_json=input_example_json)
        request_template = match_result.get("request_template", {})
        request_body_wrapper = match_result.get("request_body_wrapper", "")

        success_example = doc_info.get("success_example", {})
        output_params = doc_info.get("output_params", [])
        output_params_json = json.dumps(output_params, ensure_ascii=False)

        # Step 3: LLM 分析出参
        response_extract: dict = {}
        field_transform: dict = {}
        analysis = ""

        llm_ctx = ""
        if success_example:
            sample_str = json.dumps(success_example, ensure_ascii=False, indent=2)
            if len(sample_str) > 5000:
                sample_str = sample_str[:5000] + "\n... (已截断)"
            llm_ctx = f"接口响应样例 JSON：\n{sample_str}"
            if output_params:
                desc_lines = [
                    f"  {p.get('path', p.get('name',''))}: {p.get('desc','')}"
                    for p in output_params[:80]
                    if p.get('desc', '').strip()
                ]
                if desc_lines:
                    llm_ctx += (
                        f"\n\n接口文档出参字段说明（共 {len(output_params)} 个字段，"
                        f"请结合说明理解字段含义，优先以说明为准判断归属域）：\n"
                        + "\n".join(desc_lines)
                    )
        elif output_params:
            lines = [
                f"  {p.get('path','')}: {p.get('type','string')} — {p.get('desc','')}"
                for p in output_params[:60]
            ]
            llm_ctx = (
                f"文档未包含出参成功示例，以下是出参说明列表（共 {len(output_params)} 个字段）：\n"
                + "\n".join(lines)
            )

        if llm_ctx:
            try:
                llm_prompt = PARSE_DOCX_SYSTEM_PROMPT.replace(
                    "{llm_ctx}", f"接口响应样例：\n{llm_ctx}\n\n"
                )
                llm = _get_llm()
                raw_llm = await llm.generate(
                    llm_prompt, temperature=0.0, max_tokens=4000,
                    stage="parse_docx.llm_map", provider="parse_docx_preview",
                )
                json_m = re.search(r'\{[\s\S]*\}', raw_llm)
                if json_m:
                    mc = json.loads(json_m.group())
                    response_extract = mc.get("response_extract", {}) or {}
                    field_transform = mc.get("field_transform", {}) or {}
                    llm_analysis = mc.get("analysis", "")
                    if field_transform:
                        field_transform = _detect_units(
                            json.dumps(field_transform, ensure_ascii=False),
                            output_params_json if output_params else None,
                        )
                    analysis = "【LLM辅助生成】" + (llm_analysis or "字段映射完成")
                else:
                    logger.warning("[parse_docx_preview] LLM 返回内容无法解析为 JSON")
                    analysis = "LLM 分析完成但返回格式异常，请手动填写出参映射"
            except Exception as llm_e:
                logger.warning(f"[parse_docx_preview] LLM 辅助映射失败: {llm_e}")
                analysis = f"LLM 映射失败（{llm_e}），请手动填写或粘贴出参样例后点击「LLM 重新生成映射」"
        else:
            analysis = (
                "文档中未找到出参成功示例 JSON，且出参说明为空。\n"
                "请在下方「出参成功示例 JSON」中粘贴样例后点击「🤖 LLM 重新生成映射」。"
            )

        # 构建入参匹配展示列表
        _PHOLDER_SRC = {
            "{{PHONE}}": ("direct", "request_data.phone"),
            "{{INTENT}}": ("direct", "request_data.intent"),
            "{{CALL_ID}}": ("direct", "request_data.callId"),
            "{{TASK_ID}}": ("direct", "request_data.callId"),
            "{{PROVINCE}}": ("direct", "request_data.province"),
            "{{TOP_N}}": ("direct", "request_data.topN"),
            "{{CURRENT_OFFER_NAME}}": ("extra_data", "extra_data.currentMainOffer.curOfferName"),
            "{{CURRENT_OFFER_ID}}": ("extra_data", "extra_data.currentMainOffer.curOfferId"),
            "{{CURRENT_OFFER_FEE}}": ("extra_data", "extra_data.currentMainOffer.curOfferFee"),
        }

        def _flatten_tpl(tpl, prefix=""):
            rows = []
            for k, v in (tpl or {}).items():
                key = f"{prefix}.{k}" if prefix else k
                if isinstance(v, dict):
                    rows.extend(_flatten_tpl(v, key))
                else:
                    rows.append((key, str(v)))
            return rows

        def _resolve_match_type(placeholder: str):
            ph = str(placeholder)
            if ph in _PHOLDER_SRC:
                return _PHOLDER_SRC[ph]
            m = re.match(r'^\{\{(extra_data\..+)\}\}$', ph)
            if m:
                return ("extra_data", m.group(1))
            if "EXTRA_" in ph:
                field = ph.strip("{}").replace("EXTRA_", "").lower()
                return ("extra_data", f"extra_data.{field}")
            return ("unmatched", "需手动配置")

        # 构建入参表索引：支持多种 key 格式的查找
        # 辽宁等文档的参数名可能含前缀，如 "params.userMobile"、"Params.crmpfPubInfo.staffId"
        # 用 leaf（最末段，小写）做后备索引，支持大小写不敏感匹配
        raw_input_params = doc_info.get("input_params", [])
        _param_by_name: dict = {}      # 精确匹配
        _param_by_leaf: dict = {}      # leaf 小写匹配（后备）
        _param_by_path: dict = {}      # path 小写匹配（去除 wrapper 前缀后）
        for p in raw_input_params:
            pname = (p.get("name") or "").strip()
            if not pname:
                continue
            _param_by_name[pname] = p
            leaf = pname.split(".")[-1].lower()
            if leaf not in _param_by_leaf:
                _param_by_leaf[leaf] = p
            # 去除 wrapper 前缀后的路径（如 params.userMobile → userMobile）
            parts = pname.split(".")
            if len(parts) > 1:
                sub_path = ".".join(parts[1:]).lower()
                if sub_path not in _param_by_path:
                    _param_by_path[sub_path] = p

        def _lookup_param(field_path: str) -> dict:
            """按 field_path 查找入参表条目，支持多种前缀格式。"""
            # 1. 精确匹配
            if field_path in _param_by_name:
                return _param_by_name[field_path]
            # 2. 大小写不敏感精确匹配
            fp_lower = field_path.lower()
            for k, v in _param_by_name.items():
                if k.lower() == fp_lower:
                    return v
            # 3. wrapper 前缀去除后的路径匹配（如 request_body_wrapper 存在时）
            if field_path.lower() in _param_by_path:
                return _param_by_path[field_path.lower()]
            # 4. leaf 匹配（最末节点，小写）
            leaf = field_path.split(".")[-1].lower()
            if leaf in _param_by_leaf:
                return _param_by_leaf[leaf]
            return {}

        param_matches = []
        for field_path, placeholder in _flatten_tpl(request_template):
            pinfo = _lookup_param(field_path)
            mt, src = _resolve_match_type(placeholder)
            param_matches.append({
                "api_param": field_path,
                "type": pinfo.get("type", "string"),
                "desc": pinfo.get("desc", ""),
                "required": pinfo.get("required", True),
                "match_type": mt,
                "placeholder": str(placeholder),
                "source_path": src,
            })

        # 出参映射预览
        domain_mapping_preview: dict = {}
        try:
            ex = success_example if isinstance(success_example, dict) else {}
            if ex and (response_extract or field_transform):
                domain_mapping_preview = _apply_mapping(ex, response_extract or {}, field_transform or {})
        except Exception as dprev_e:
            logger.warning(f"[parse_docx_preview] domain_mapping_preview 失败: {dprev_e}")

        return {
            "code": 0,
            "message": "解析完成",
            "data": {
                "basic_info": {
                    "api_name": doc_info.get("api_name", ""),
                    "description": doc_info.get("description", ""),
                    "url": doc_info.get("url", ""),
                    "method": doc_info.get("method", "POST"),
                    "headers": doc_info.get("headers", {}),
                    "version": doc_info.get("version", ""),
                },
                "param_matches": param_matches,
                "request_body_wrapper": request_body_wrapper,
                "request_template": request_template,
                "response_extract": response_extract,
                "field_transform": field_transform,
                "analysis": analysis,
                "success_example": success_example,
                "domain_mapping_preview": domain_mapping_preview,
                "unit_conversions": field_transform.get("_unit_conversions", []),
                "extra_data_example": {
                    "currentMainOffer": {
                        "curOfferName": "128元5G套餐",
                        "curOfferId": "111601000461",
                        "curOfferFee": "12800",
                    }
                },
                "input_params": doc_info.get("input_params", []),
                "output_params": doc_info.get("output_params", []),
            },
        }
    except Exception as e:
        logger.error(f"[parse_docx_preview] 失败: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/api/interfaces/apply_parsed")
async def apply_parsed_interface(body: Dict[str, Any]):
    """将解析/编辑后的接口配置直接写入 api_nodes.json 并热重载。

    body: province, intent, api_name, url, method, description, headers,
          request_template, response_extract, field_transform,
          mock_response, success_example, mock_mode
    """
    import datetime

    province = body.get("province", "")
    intent = body.get("intent", "")
    api_name = body.get("api_name", "")

    if not (province and intent and api_name):
        raise HTTPException(400, "province / intent / api_name 必填")

    skill_registry = _get_skill_registry()
    pkg = skill_registry.get(province, intent)
    if pkg is None:
        raise HTTPException(404, f"技能包不存在: {province}/{intent}")

    # mock_response 处理：优先非空的 mock_response；空对象时用 success_example 补全
    _mr = body.get("mock_response")
    _se = body.get("success_example")
    if isinstance(_mr, dict) and len(_mr) == 0:
        mock_response = _se if _se else _mr
    elif _mr is not None:
        mock_response = _mr
    else:
        mock_response = _se

    mock_mode_in = body.get("mock_mode")
    default_mock_mode = mock_mode_in if isinstance(mock_mode_in, bool) else bool(mock_response)

    new_node: dict = {
        "enabled": True,
        "_comment": body.get("description", ""),
        "url": body.get("url", ""),
        "method": body.get("method", "POST"),
        "headers": body.get("headers") or {},
        "timeout": 30,
        "max_retries": 2,
        "mock_mode": default_mock_mode,
        "request_template": body.get("request_template") or {},
        "response_extract": body.get("response_extract") or {},
        "field_transform": body.get("field_transform") or {},
        "created_by": "interface_mapper_agent",
        "created_at": datetime.datetime.now().strftime("%Y-%m-%d %H:%M"),
    }
    # 仅当 request_body_wrapper 非空时才写入
    _rbw = body.get("request_body_wrapper", "")
    if _rbw:
        new_node["request_body_wrapper"] = _rbw
    if mock_response:
        new_node["mock_response"] = mock_response

    api_nodes_cfg = dict(pkg.config.get("api_nodes", {}))

    # 保留元注释字段（_comment、_desc 等）
    meta_keys = {k: v for k, v in api_nodes_cfg.items() if k.startswith("_")}
    other_nodes = {k: v for k, v in api_nodes_cfg.items() if not k.startswith("_") and k != api_name}
    api_nodes_cfg = {**meta_keys, api_name: new_node, **other_nodes}

    ok = skill_registry.save_api_nodes(province, intent, api_nodes_cfg)
    if not ok:
        raise HTTPException(500, "写入 api_nodes.json 失败")

    logger.info(f"[apply_parsed] 接口 {province}/{intent}/{api_name} 配置已写入")

    # ── 同步推断 field_aliases 并写入 biz_config.json ────────────────────
    try:
        from management.interface_mapper.scripts.tools import (
            _infer_field_aliases_from_api_nodes,
            _infer_field_aliases_from_output_params,
            _merge_field_aliases,
        )
        output_params = body.get("output_params") or []

        # 路1：从 api_nodes 套餐域字段 + output_params 说明推断（最精准）
        inferred_from_nodes = _infer_field_aliases_from_api_nodes(
            {api_name: new_node}, output_params
        )

        # 路2：从 output_params 路径说明推断
        inferred_from_params = {}
        response_extract_data = body.get("response_extract") or {}
        if output_params and response_extract_data:
            inferred_from_params = _infer_field_aliases_from_output_params(
                output_params, response_extract_data
            )

        # 合并两路（api_nodes 推断优先）
        combined: dict = dict(inferred_from_params)
        for sem_key, fields in inferred_from_nodes.items():
            if sem_key not in combined:
                combined[sem_key] = []
            for f in fields:
                if f not in combined[sem_key]:
                    combined[sem_key].insert(0, f)

        if combined:
            # 读取现有 biz_config
            existing_biz = pkg.config.get("biz_config") or {}
            existing_aliases = existing_biz.get("field_aliases") or {}
            # 过滤掉注释键
            explicit = {k: v for k, v in existing_aliases.items() if not k.startswith("_")}

            final_aliases = _merge_field_aliases(combined, explicit or None)
            new_biz = dict(existing_biz)
            new_biz["field_aliases"] = {
                "_comment": "套餐字段别名（按优先级列出，省份专属字段名在前）",
                **final_aliases,
            }
            ok_biz = skill_registry.save_biz_config(province, intent, new_biz)
            if ok_biz:
                logger.info(
                    f"[apply_parsed] biz_config.json field_aliases 已更新: {list(final_aliases.keys())}"
                )
            else:
                logger.warning(f"[apply_parsed] biz_config.json 更新失败")
        else:
            logger.debug(f"[apply_parsed] 未推断出 field_aliases，biz_config 不更新")
    except Exception as alias_err:
        logger.warning(f"[apply_parsed] field_aliases 推断失败（忽略，不影响接口写入）: {alias_err}")

    return {
        "code": 0,
        "message": f"接口 {api_name} 配置已写入 api_nodes.json，配置生效",
        "data": {"api_name": api_name, "province": province, "intent": intent},
    }


@router.post("/api/skills/preview_mapping")
async def preview_mapping(req: PreviewMappingRequest):
    """用 mock_response 跑实际映射函数，返回各数据域的真实结果供前端展示。"""
    try:
        domain_result = _apply_mapping(req.mock_response, req.response_extract, req.field_transform)
        return {"code": 0, "data": {"domain_result": domain_result}}
    except Exception as e:
        logger.error(f"[preview_mapping] 失败: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/api/skills/refine_mapping_preview")
async def refine_mapping_preview(req: RefineMappingPreviewRequest):
    """根据用户修改后的数据域结果反推 extract/transform，仅返回预览，不写 api_nodes.json。"""
    try:
        mock_resp = req.mock_response or {}
        new_extract, new_transform, analysis = await _run_refine_mapping_llm(
            mock_resp,
            req.response_extract or {},
            req.field_transform or {},
            req.user_domain_result,
            llm_stage="interface_mapper.refine_preview",
        )
        domain_result = _apply_mapping(mock_resp, new_extract, new_transform)
        return {
            "code": 0,
            "data": {
                "response_extract": new_extract,
                "field_transform": new_transform,
                "domain_result": domain_result,
                "analysis": analysis,
            },
        }
    except Exception as e:
        logger.error(f"[refine_mapping_preview] 失败: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/api/skills/refine_mapping")
async def refine_mapping(req: RefineMappingRequest):
    """用户修改数据域结果后，LLM 反推并更新 api_nodes.json 的映射规则。"""
    try:
        api_nodes = req.current_api_nodes
        node_key = next((k for k in api_nodes if not k.startswith("_")), None)
        if not node_key:
            raise HTTPException(status_code=400, detail="api_nodes 中未找到接口节点")
        node = api_nodes[node_key]

        current_extract = node.get("response_extract", {})
        current_transform = node.get("field_transform", {})
        mock_resp = req.mock_response or node.get("mock_response", {})

        new_extract, new_transform, analysis = await _run_refine_mapping_llm(
            mock_resp, current_extract, current_transform, req.user_domain_result,
            llm_stage="interface_mapper.refine",
        )

        updated_nodes = json.loads(json.dumps(api_nodes))
        updated_nodes[node_key]["response_extract"] = new_extract
        updated_nodes[node_key]["field_transform"] = new_transform

        # 写回文件
        nodes_path = Path("skills-runtime") / req.province / req.intent / "config" / "api_nodes.json"
        if nodes_path.exists():
            nodes_path.write_text(
                json.dumps(updated_nodes, ensure_ascii=False, indent=2),
                encoding="utf-8",
            )

        domain_result = _apply_mapping(mock_resp, new_extract, new_transform)
        return {
            "code": 0,
            "data": {
                "response_extract": new_extract,
                "field_transform": new_transform,
                "domain_result": domain_result,
                "analysis": analysis,
                "api_nodes": updated_nodes,
            },
        }
    except Exception as e:
        logger.error(f"[refine_mapping] 失败: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/api/skills/create_from_doc")
async def create_skill_from_doc(req: CreateSkillFromDocRequest):
    """上传接口文档（base64），通过 Interface Mapper Agent 自动解析并生成 Skill 包。"""
    try:
        from management.interface_mapper.scripts.agent_runner import run as agent_run
        result = await agent_run(
            docx_content_b64=req.docx_content_b64,
            province=req.province,
            intent=req.intent,
        )
        if result.get("skill_path"):
            try:
                _get_skill_registry().reload()
            except Exception:
                pass
        return {"code": 0, "message": "Skill 包生成成功", "data": result}
    except Exception as e:
        logger.error(f"[create_skill_from_doc] 失败: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/api/download/interface_template")
async def download_interface_template():
    """下载接口规范文档模板（动态生成 docx）。"""
    try:
        from docx import Document
        from fastapi.responses import StreamingResponse

        doc = Document()
        doc.add_heading("接口规范文档模板", 0)
        doc.add_paragraph("此文档用于「上传文档自主解析」功能。AI 将自动提取以下信息生成 api_nodes.json 配置。")
        doc.add_heading("文档结构要求（必须包含以下部分）", level=1)
        for line in [
            "1. 接口名称：营销推荐接口",
            "2. 接口描述：获取用户当前套餐、推荐套餐及用量标签",
            "3. 请求地址：http://xxx/znhs/marketing/recommend",
            "4. 请求方式：POST",
            "5. 请求头：Content-Type: application/json",
            "6. 接口版本：V1.0",
        ]:
            doc.add_paragraph(line)

        doc.add_heading("入参说明", level=2)
        table = doc.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        for i, h in enumerate(["参数名", "类型", "是否必选", "说明", "示例"]):
            table.rows[0].cells[i].text = h
        for params in [
            ("phone", "string", "是", "用户手机号", "13800138000"),
            ("intent", "string", "是", "业务意图", "套餐推荐"),
            ("ioId", "string", "否", "会话ID", "io_001"),
        ]:
            row_cells = table.add_row().cells
            for i, val in enumerate(params):
                row_cells[i].text = val

        doc.add_heading("入参成功示例", level=2)
        doc.add_paragraph('{\n  "phone": "13800138000",\n  "intent": "套餐推荐"\n}')
        doc.add_heading("出参说明", level=2)
        table2 = doc.add_table(rows=1, cols=4)
        table2.style = "Table Grid"
        for i, h in enumerate(["参数路径", "类型", "是否必返", "说明"]):
            table2.rows[0].cells[i].text = h
        doc.add_heading("出参成功示例", level=2)
        doc.add_paragraph('{\n  "rtnCode": "0",\n  "rtnMsg": "成功",\n  "bean": {\n    "mainoffer": {...}\n  }\n}')
        doc.add_heading("出参失败示例", level=2)
        doc.add_paragraph('{\n  "rtnCode": "1",\n  "rtnMsg": "失败"\n}')
        doc.add_paragraph("\n保存为 .docx 后上传即可由 Agent 自动解析生成接口配置和映射规则。")

        bio = BytesIO()
        doc.save(bio)
        bio.seek(0)
        return StreamingResponse(
            bio,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=接口规范文档模板.docx"},
        )
    except ImportError:
        raise HTTPException(500, "python-docx 未安装，无法生成模板")
    except Exception as e:
        logger.error(f"生成模板失败: {e}")
        raise HTTPException(500, f"模板生成失败: {e}")
